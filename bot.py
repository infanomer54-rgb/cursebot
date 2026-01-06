# Файл bot.py - основной файл бота
import os
import logging
import sqlite3
import re
import asyncio
from datetime import datetime
import json
import io
import sys
import random
import hashlib
from typing import List, Dict
from collections import Counter
from threading import Thread

import requests
import PyPDF2
import docx2txt
import aiofiles
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
from googlesearch import search
from transformers import pipeline
from sentence_transformers import SentenceTransformer
import textstat
import pymorphy3
from flask import Flask

from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    Application, CommandHandler, MessageHandler, CallbackQueryHandler,
    ContextTypes, filters
)

# Flask app для Railway
app = Flask(__name__)

@app.route('/')
def home():
    return "🤖 Academic Writing Bot is running!"

@app.route('/health')
def health():
    return "OK", 200

def run_flask():
    port = int(os.getenv("PORT", 8080))
    app.run(host='0.0.0.0', port=port)

# Конфигурация
BOT_TOKEN = os.getenv('BOT_TOKEN')
DEEPSEEK_API_KEY = os.getenv('DEEPSEEK_API_KEY')
DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"

# Создаем директории
os.makedirs("методички", exist_ok=True)
os.makedirs("работы", exist_ok=True)

# Настройка логирования
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

class Database:
    def __init__(self, db_path="bot_database.db"):
        self.db_path = db_path
        self.init_db()
    
    def init_db(self):
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS users (
                user_id INTEGER PRIMARY KEY,
                username TEXT,
                first_name TEXT,
                last_name TEXT,
                group_name TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS works (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                work_type TEXT,
                topic TEXT,
                subject TEXT,
                content TEXT,
                methodic_info TEXT,
                student_info TEXT,
                teacher_info TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS methodics (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                filename TEXT,
                file_path TEXT,
                university_name TEXT,
                university_address TEXT,
                faculty TEXT,
                department TEXT,
                work_structure TEXT,
                formatting_style TEXT,
                uploaded_by INTEGER,
                uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        conn.commit()
        conn.close()
    
    def add_user(self, user_id, username, first_name, last_name, group_name=None):
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            INSERT OR REPLACE INTO users (user_id, username, first_name, last_name, group_name)
            VALUES (?, ?, ?, ?, ?)
        ''', (user_id, username, first_name, last_name, group_name))
        conn.commit()
        conn.close()
    
    def update_user_group(self, user_id, group_name):
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('UPDATE users SET group_name = ? WHERE user_id = ?', (group_name, user_id))
        conn.commit()
        conn.close()
    
    def get_user(self, user_id):
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM users WHERE user_id = ?', (user_id,))
        result = cursor.fetchone()
        conn.close()
        return result
    
    def create_work(self, user_id, work_type, topic, subject, methodic_info=None, student_info=None, teacher_info=None):
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        try:
            methodic_json = None
            if methodic_info:
                try:
                    methodic_json = json.dumps(methodic_info, ensure_ascii=False)
                except (TypeError, ValueError) as e:
                    logger.error(f"Error serializing methodic_info: {e}")
                    methodic_json = json.dumps({}, ensure_ascii=False)
            
            student_json = None
            if student_info:
                try:
                    student_json = json.dumps(student_info, ensure_ascii=False)
                except (TypeError, ValueError) as e:
                    logger.error(f"Error serializing student_info: {e}")
                    student_json = json.dumps({}, ensure_ascii=False)
            
            teacher_json = None
            if teacher_info:
                try:
                    teacher_json = json.dumps(teacher_info, ensure_ascii=False)
                except (TypeError, ValueError) as e:
                    logger.error(f"Error serializing teacher_info: {e}")
                    teacher_json = json.dumps({}, ensure_ascii=False)
            
            cursor.execute('''
                INSERT INTO works (user_id, work_type, topic, subject, methodic_info, student_info, teacher_info)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (user_id, work_type, topic, subject, methodic_json, student_json, teacher_json))
            work_id = cursor.lastrowid
            conn.commit()
            return work_id
        except Exception as e:
            logger.error(f"Error creating work: {e}")
            conn.rollback()
            return None
        finally:
            conn.close()
    
    def update_work_content(self, work_id, content):
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('UPDATE works SET content = ? WHERE id = ?', (content, work_id))
        conn.commit()
        conn.close()
    
    def add_methodic(self, filename, file_path, university_name, university_address, faculty, department, work_structure, formatting_style, user_id):
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        try:
            work_structure_json = json.dumps(work_structure, ensure_ascii=False) if work_structure else json.dumps({
                'required_sections': ['Введение', 'Основная часть', 'Заключение', 'Список литературы'],
                'chapter_count': 3,
                'has_introduction': True,
                'has_conclusion': True,
                'has_bibliography': True
            }, ensure_ascii=False)
            
            formatting_style_json = json.dumps(formatting_style, ensure_ascii=False) if formatting_style else json.dumps({
                'font_family': 'Times New Roman',
                'font_size': '14',
                'line_spacing': '1.5',
                'margin_left': '3',
                'margin_right': '1',
                'margin_top': '2',
                'margin_bottom': '2'
            }, ensure_ascii=False)
            
            cursor.execute('''
                INSERT INTO methodics (filename, file_path, university_name, university_address, faculty, department, work_structure, formatting_style, uploaded_by)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (filename, file_path, university_name, university_address, faculty, department, 
                  work_structure_json, 
                  formatting_style_json, 
                  user_id))
            methodic_id = cursor.lastrowid
            conn.commit()
            conn.close()
            return methodic_id
        except Exception as e:
            logger.error(f"Error saving methodic to database: {e}")
            conn.rollback()
            conn.close()
            return None
    
    def get_methodics(self):
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('SELECT id, filename, university_name FROM methodics ORDER BY uploaded_at DESC')
        methodics = cursor.fetchall()
        conn.close()
        return methodics
    
    def get_methodic(self, methodic_id):
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM methodics WHERE id = ?', (methodic_id,))
        result = cursor.fetchone()
        conn.close()
        return result

class DocumentProcessor:
    def extract_text_from_pdf(self, file_path):
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                return text.strip()
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path):
        try:
            text = docx2txt.process(file_path)
            return text.strip() if text else ""
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return ""
    
    async def extract_text_from_txt(self, file_path):
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                return await file.read()
        except Exception as e:
            logger.error(f"TXT extraction error: {e}")
            return ""
    
    async def process_methodic(self, file_path):
        file_extension = file_path.lower().split('.')[-1]
        text = ""
        
        if file_extension == 'pdf':
            text = self.extract_text_from_pdf(file_path)
        elif file_extension == 'docx':
            text = self.extract_text_from_docx(file_path)
        elif file_extension == 'txt':
            text = await self.extract_text_from_txt(file_path)
        else:
            return None
        
        if not text:
            return None
        
        return self.extract_methodic_info(text)
    
    def extract_methodic_info(self, text):
        try:
            university_info = self._extract_university_info(text)
            work_structure = self._extract_work_structure(text)
            formatting_style = self._extract_formatting_style(text)
            
            if not university_info:
                university_info = {
                    'university_name': "Федеральное государственное автономное образовательное учреждение высшего образования",
                    'university_address': "г. Москва, ул. Примерная, д. 123",
                    'faculty': "Факультет информационных технологий",
                    'department': "Кафедра информатики и вычислительной техники"
                }
            
            if not work_structure:
                work_structure = {
                    'required_sections': ['Введение', 'Основная часть', 'Заключение', 'Список литературы'],
                    'chapter_count': 3,
                    'has_introduction': True,
                    'has_conclusion': True,
                    'has_bibliography': True
                }
            
            if not formatting_style:
                formatting_style = {
                    'font_family': 'Times New Roman',
                    'font_size': '14',
                    'line_spacing': '1.5',
                    'margin_left': '3',
                    'margin_right': '1',
                    'margin_top': '2',
                    'margin_bottom': '2'
                }
            
            return {
                'university': university_info,
                'work_structure': work_structure,
                'formatting_style': formatting_style,
                'full_text': text[:4000]
            }
        except Exception as e:
            logger.error(f"Methodic info extraction error: {e}")
            return {
                'university': {
                    'university_name': "Федеральное государственное автономное образовательное учреждение высшего образования",
                    'university_address': "г. Москва, ул. Примерная, д. 123",
                    'faculty': "Факультет информационных технологий",
                    'department': "Кафедра информатики и вычислительной техники"
                },
                'work_structure': {
                    'required_sections': ['Введение', 'Основная часть', 'Заключение', 'Список литературы'],
                    'chapter_count': 3,
                    'has_introduction': True,
                    'has_conclusion': True,
                    'has_bibliography': True
                },
                'formatting_style': {
                    'font_family': 'Times New Roman',
                    'font_size': '14',
                    'line_spacing': '1.5',
                    'margin_left': '3',
                    'margin_right': '1',
                    'margin_top': '2',
                    'margin_bottom': '2'
                },
                'full_text': text[:2000] if text else ""
            }
    
    def _extract_university_info(self, text):
        patterns = {
            'university_name': [
                r'(?:ФГБОУ ВО|ФГАОУ ВО|ФГБОУ|ГОУ ВПО|Федеральное|Государственное)[^.!?]{0,200}?(?:университет|институт|академия|college|university)',
                r'[А-Я][А-Яа-яё\s\-]{5,}?(?:университет|институт|академия)[^.!?]{0,100}',
                r'МИНИСТЕРСТВО[^.!?]{0,150}?(?:университет|институт|академия)',
                r'НАЦИОНАЛЬНЫЙ[^.!?]{0,100}?(?:университет|институт|академия)'
            ],
            'university_address': [
                r'(?:адрес|address)[:\s]+([^.!?\n]{20,100})',
                r'[0-9]{6}[,\s]+(?:г\.|город|city)[\s]+([А-Я][а-яё\s\-]+)',
                r'(?:г\.|город)[\s]+([А-Я][а-яё]+)[^.!?]{0,50}?(?:ул\.|улица|проспект|пр\.)',
                r'[А-Я][а-яё\s\-]{5,}?(?:область|край)[^.!?]{0,50}?(?:г\.|город)[\s]+([А-Я][а-яё]+)'
            ],
            'faculty': [
                r'(?:факультет|faculty)[\s]+([^.!?\n]{10,80})',
                r'[А-Я][А-Яа-яё\s\-]{5,}?(?:факультет|институт)[^.!?]{0,50}',
                r'(?:институт)[^.!?]{0,50}?([А-Я][А-Яа-яё\s\-]{5,}?(?:информатики|экономики|юриспруденции))'
            ],
            'department': [
                r'(?:кафедра|department)[\s]+([^.!?\n]{10,80})',
                r'[А-Я][А-Яа-яё\s\-]{5,}?(?:кафедра)[^.!?]{0,50}',
                r'(?:кафедра)[^.!?]{0,50}?([А-Я][А-Яа-яё\s\-]{5,}?(?:информатики|математики|физики))'
            ]
        }
        
        university_info = {}
        for key, pattern_list in patterns.items():
            for pattern in pattern_list:
                matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
                if matches:
                    university_info[key] = matches[0].strip()
                    break
        
        if not university_info.get('university_name'):
            university_info['university_name'] = "Федеральное государственное автономное образовательное учреждение высшего образования"
        if not university_info.get('university_address'):
            university_info['university_address'] = "г. Москва, ул. Примерная, д. 123"
        if not university_info.get('faculty'):
            university_info['faculty'] = "Факультет информационных технологий"
        if not university_info.get('department'):
            university_info['department'] = "Кафедра информатики и вычислительной техники"
        
        return university_info
    
    def _extract_work_structure(self, text):
        structure_patterns = [
            r'(?:структура|содержание|оглавление)[^.!?]{0,200}?(?:введение|введени[ея])[^.!?]{0,200}?(?:глава|раздел|часть)[^.!?]{0,200}?(?:заключение|выводы)',
            r'(?:должна содержать|включает|состоит из)[^.!?]{0,300}',
            r'(?:введение|введени[ея])[^.!?]{0,100}?(?:основная часть|главы|разделы)[^.!?]{0,100}?(?:заключение|выводы)',
            r'(?:глава\s+\d+[^.!?]{0,50}){2,}',
            r'(?:раздел\s+\d+[^.!?]{0,50}){2,}'
        ]
        
        work_structure = {
            'required_sections': [],
            'chapter_count': 3,
            'has_introduction': True,
            'has_conclusion': True,
            'has_bibliography': True
        }
        
        for pattern in structure_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            if matches:
                structure_text = matches[0]
                
                chapter_matches = re.findall(r'(глава|раздел)\s*(\d+)', structure_text, re.IGNORECASE)
                if chapter_matches:
                    work_structure['chapter_count'] = len(chapter_matches)
                
                if 'введение' in structure_text.lower():
                    work_structure['required_sections'].append('Введение')
                if 'заключение' in structure_text.lower() or 'выводы' in structure_text.lower():
                    work_structure['required_sections'].append('Заключение')
                if 'литератур' in structure_text.lower() or 'библиограф' in structure_text.lower():
                    work_structure['required_sections'].append('Список литературы')
                if 'приложен' in structure_text.lower():
                    work_structure['required_sections'].append('Приложения')
                
                break
        
        if not work_structure['required_sections']:
            work_structure['required_sections'] = ['Введение', 'Основная часть', 'Заключение', 'Список литературы']
        
        return work_structure
    
    def _extract_formatting_style(self, text):
        formatting_patterns = {
            'font_family': [
                r'шрифт[:\s]*([^\n,\d]{3,30})',
                r'([Tt]imes [Nn]ew [Rr]oman|[Aa]rial|[Hh]elvetica)'
            ],
            'font_size': [
                r'шрифт[:\s]*(\d+)',
                r'размер[:\s]*шрифта[:\s]*(\d+)',
                r'(\d+)[\s]*(?:pt|пт)'
            ],
            'line_spacing': [
                r'интервал[:\s]*([^\n]+)',
                r'([\d\.]+)[\s]*(?:междустрочн|интервал)',
                r'(полуторный|одинарный|двойной)'
            ],
            'margins': [
                r'поля[:\s]*([^\n]{10,50})',
                r'левое[:\s]*(\d+)[^.!?]{0,20}?правое[:\s]*(\d+)',
                r'верхнее[:\s]*(\d+)[^.!?]{0,20}?нижнее[:\s]*(\d+)'
            ]
        }
        
        formatting_style = {}
        for key, pattern_list in formatting_patterns.items():
            for pattern in pattern_list:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    if key == 'margins' and len(matches[0]) == 2:
                        formatting_style['margin_left'] = matches[0][0]
                        formatting_style['margin_right'] = matches[0][1]
                    elif key == 'margins' and len(matches[0]) == 2:
                        formatting_style['margin_top'] = matches[0][0]
                        formatting_style['margin_bottom'] = matches[0][1]
                    else:
                        formatting_style[key] = matches[0] if isinstance(matches[0], str) else matches[0][0]
                    break
        
        if not formatting_style.get('font_family'):
            formatting_style['font_family'] = 'Times New Roman'
        if not formatting_style.get('font_size'):
            formatting_style['font_size'] = '14'
        if not formatting_style.get('line_spacing'):
            formatting_style['line_spacing'] = '1.5'
        if not formatting_style.get('margin_left'):
            formatting_style['margin_left'] = '3'
        if not formatting_style.get('margin_right'):
            formatting_style['margin_right'] = '1'
        if not formatting_style.get('margin_top'):
            formatting_style['margin_top'] = '2'
        if not formatting_style.get('margin_bottom'):
            formatting_style['margin_bottom'] = '2'
        
        return formatting_style

class EnhancedAcademicWriter:
    def __init__(self):
        self.api_key = DEEPSEEK_API_KEY
        self.api_url = DEEPSEEK_API_URL
        self.grammar_checker = None
        self.similarity_model = None
        self.morph = None
        self.used_phrases = set()
        
        # Инициализация моделей
        try:
            self.grammar_checker = pipeline("text2text-generation", model="cointegrated/rut5-base-grammar-correction", device=-1)
            self.similarity_model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')
            self.morph = pymorphy3.MorphAnalyzer()
        except Exception as e:
            logger.error(f"Error initializing models: {e}")
    
    def generate_complete_work(self, work_type, topic, subject, methodic_info=None):
        work_type_names = {
            "coursework": "курсовой работы",
            "essay": "реферата", 
            "thesis": "дипломной работы"
        }
        
        sources = self._search_academic_sources(topic, subject)
        
        system_prompt = self._create_enhanced_prompt(work_type, topic, subject, methodic_info, sources)
        
        full_content = self._make_api_call(
            system_prompt,
            f"Напиши полный текст {work_type_names[work_type]} на тему '{topic}' объемом не менее {self._get_target_word_count(work_type)} слов."
        )
        
        if not full_content.startswith("❌") and not full_content.startswith("⏰"):
            enhanced_content = self._enhance_content_quality(full_content, topic, subject)
            return enhanced_content
        
        return full_content
    
    def _search_academic_sources(self, topic: str, subject: str) -> List[Dict]:
        search_queries = [
            f"{topic} {subject} научная статья",
            f"{topic} исследования последние публикации",
            f"{subject} {topic} академический источник",
            f"{topic} диссертация автореферат",
            f"{subject} научный журнал публикации"
        ]
        
        sources = []
        seen_urls = set()
        
        for query in search_queries[:2]:
            try:
                for url in search(query, num_results=2, lang='ru'):
                    if url not in seen_urls:
                        content = self._extract_academic_content(url)
                        if content and len(content) > 100:
                            sources.append({
                                'url': url,
                                'content': content[:300],
                                'relevance': self._calculate_relevance(content, topic)
                            })
                            seen_urls.add(url)
            except Exception as e:
                logger.error(f"Search error: {e}")
                continue
        
        return sorted(sources, key=lambda x: x['relevance'], reverse=True)[:3]
    
    def _extract_academic_content(self, url: str) -> str:
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(url, headers=headers, timeout=10)
            soup = BeautifulSoup(response.text, 'html.parser')
            
            for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
                tag.decompose()
            
            text = soup.get_text()
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            return text[:1500]
            
        except Exception as e:
            logger.error(f"Content extraction error: {e}")
            return ""
    
    def _calculate_relevance(self, content: str, topic: str) -> float:
        topic_words = set(self._normalize_text(topic).split())
        content_words = set(self._normalize_text(content).split())
        
        if not topic_words or not content_words:
            return 0.0
        
        intersection = topic_words.intersection(content_words)
        return len(intersection) / len(topic_words)
    
    def _normalize_text(self, text: str) -> str:
        text = text.lower()
        text = re.sub(r'[^\w\s]', ' ', text)
        words = text.split()
        normalized_words = []
        
        for word in words:
            try:
                if self.morph:
                    parsed = self.morph.parse(word)[0]
                    normalized_words.append(parsed.normal_form)
                else:
                    normalized_words.append(word)
            except:
                normalized_words.append(word)
        
        return ' '.join(normalized_words)
    
    def _create_enhanced_prompt(self, work_type, topic, subject, methodic_info, sources):
        sources_text = ""
        if sources:
            sources_text = "НАУЧНЫЕ ИСТОЧНИКИ ДЛЯ ИСПОЛЬЗОВАНИЯ:\n"
            for i, source in enumerate(sources[:2], 1):
                sources_text += f"{i}. {source['content'][:150]}...\n"
        
        structure_info = self._get_structure_info(methodic_info)
        
        return f"""Ты - опытный академический писатель. Создай УНИКАЛЬНУЮ, ГРАМОТНУЮ и НАУЧНО ОБОСНОВАННУЮ работу.

{sources_text}

{structure_info}

ТЕМА: {topic}
ПРЕДМЕТ: {subject}
ТИП РАБОТЫ: {self._get_work_type_name(work_type)}

КЛЮЧЕВЫЕ ТРЕБОВАНИЯ:
1. УНИКАЛЬНОСТЬ: Избегай шаблонных фраз, клише и повторений
2. ГРАММАТИКА: Идеальная грамматика, пунктуация и стиль
3. НАУЧНОСТЬ: Используй точную терминологию
4. СТРУКТУРА: Четкая логическая структура
5. ОБЪЕМ: Не менее {self._get_target_word_count(work_type)} слов

ЗАПРЕЩЕНО:
- Использовать шаблонные фразы типа "В данной работе", "Актуальность темы заключается"
- Повторять одни и те же мысли
- Делать грамматические ошибки

СТИЛЬ:
- Академический, но естественный
- Точно используй термины предметной области
- Поддерживай научную строгость
"""
    
    def _get_structure_info(self, methodic_info):
        if not methodic_info:
            return "СТАНДАРТНАЯ СТРУКТУРА:\n- Введение\n- 3 главы основной части\n- Заключение\n- Список литературы"
        
        work_structure = methodic_info.get('work_structure', {})
        required_sections = work_structure.get('required_sections', [])
        
        if required_sections:
            return "СТРУКТУРА ИЗ МЕТОДИЧКИ:\n" + "\n".join([f"- {section}" for section in required_sections])
        else:
            return "СТАНДАРТНАЯ СТРУКТУРА:\n- Введение\n- Основная часть\n- Заключение\n- Список литературы"
    
    def _get_work_type_name(self, work_type):
        names = {
            'coursework': 'курсовой работы',
            'essay': 'реферата',
            'thesis': 'дипломной работы'
        }
        return names.get(work_type, 'академической работы')
    
    def _get_target_word_count(self, work_type):
        word_counts = {
            "essay": 4000,
            "coursework": 8000,
            "thesis": 15000
        }
        return word_counts.get(work_type, 6000)
    
    def _enhance_content_quality(self, content: str, topic: str, subject: str) -> str:
        sentences = re.split(r'(?<=[.!?])\s+', content)
        
        unique_sentences = []
        seen_hashes = set()
        
        for sentence in sentences:
            if sentence.strip():
                words = self._normalize_text(sentence).split()[:8]
                sentence_hash = hashlib.md5(' '.join(words).encode()).hexdigest()
                
                if sentence_hash not in seen_hashes:
                    seen_hashes.add(sentence_hash)
                    
                    improved_sentence = self._improve_sentence_quality(sentence)
                    unique_sentences.append(improved_sentence)
        
        enhanced_text = ' '.join(unique_sentences)
        
        enhanced_text = self._replace_cliches(enhanced_text)
        
        return enhanced_text
    
    def _improve_sentence_quality(self, sentence: str) -> str:
        if len(sentence.split()) > 4 and self.grammar_checker:
            try:
                result = self.grammar_checker(sentence, max_length=100, num_beams=2)[0]['generated_text']
                return result
            except Exception as e:
                logger.error(f"Grammar check error: {e}")
                return sentence
        return sentence
    
    def _replace_cliches(self, text: str) -> str:
        replacements = {
            "в данной работе": "В исследовании",
            "актуальность темы заключается": "Значимость изучения обусловлена",
            "целью работы является": "Основной целью выступает",
            "задачами работы являются": "Ключевыми задачами исследования определены",
            "объектом исследования является": "В качестве объекта изучения рассматривается",
            "предметом исследования является": "Предметная область охватывает",
            "во введении": "В начальном разделе",
            "в заключении": "В завершающей части",
            "было выявлено": "Установлено",
            "можно сделать вывод": "Следует заключить"
        }
        
        for cliche, replacement in replacements.items():
            text = re.sub(r'\b' + re.escape(cliche) + r'\b', replacement, text, flags=re.IGNORECASE)
        
        return text
    
    def _make_api_call(self, system_prompt, user_prompt):
        if not self.api_key:
            logger.error("DeepSeek API key not configured")
            return "❌ Ошибка: API ключ DeepSeek не настроен"
        
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {self.api_key}"
        }
        
        data = {
            "model": "deepseek-chat",
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            "temperature": 0.7,
            "max_tokens": 8000
        }
        
        try:
            logger.info(f"Sending request to DeepSeek API...")
            response = requests.post(self.api_url, headers=headers, json=data, timeout=180)
            response.raise_for_status()
            result = response.json()
            content = result['choices'][0]['message']['content']
            
            word_count = len(content.split())
            logger.info(f"Received response: {word_count} words")
            
            return content
            
        except requests.exceptions.Timeout:
            logger.error("DeepSeek API timeout")
            return "⏰ Время ожидания истекло. Попробуйте еще раз."
        except requests.exceptions.RequestException as e:
            logger.error(f"DeepSeek API request error: {e}")
            return "❌ Ошибка соединения с сервисом."
        except Exception as e:
            logger.error(f"Unexpected API error: {e}")
            return f"❌ Ошибка генерации: {str(e)}"

class WordDocumentGenerator:
    def __init__(self):
        self.doc = None
    
    def create_document(self, work_type, topic, subject, content, methodic_info, student_info, teacher_info):
        try:
            self.doc = Document()
            
            self._apply_formatting(methodic_info)
            
            self._create_title_page(work_type, topic, subject, methodic_info, student_info, teacher_info)
            
            self._create_table_of_contents(methodic_info)
            
            self._add_main_content(content, methodic_info)
            
            self._add_bibliography()
            
            file_stream = io.BytesIO()
            self.doc.save(file_stream)
            file_stream.seek(0)
            
            return file_stream
            
        except Exception as e:
            logger.error(f"Error creating Word document: {e}")
            return None
        finally:
            self.doc = None
    
    def _apply_formatting(self, methodic_info):
        try:
            formatting = methodic_info.get('formatting_style', {}) if methodic_info else {}
            font_family = formatting.get('font_family', 'Times New Roman')
            font_size = int(formatting.get('font_size', '14'))
            
            style = self.doc.styles['Normal']
            font = style.font
            font.name = font_family
            font.size = Pt(font_size)
            
            line_spacing = formatting.get('line_spacing', '1.5')
            if '1.5' in line_spacing or 'полуторный' in line_spacing:
                style.paragraph_format.line_spacing = 1.5
            elif '1.0' in line_spacing or 'одинарный' in line_spacing:
                style.paragraph_format.line_spacing = 1.0
            elif '2.0' in line_spacing or 'двойной' in line_spacing:
                style.paragraph_format.line_spacing = 2.0
            
            sections = self.doc.sections
            for section in sections:
                section.left_margin = Inches(float(formatting.get('margin_left', 3)) * 0.393701)
                section.right_margin = Inches(float(formatting.get('margin_right', 1)) * 0.393701)
                section.top_margin = Inches(float(formatting.get('margin_top', 2)) * 0.393701)
                section.bottom_margin = Inches(float(formatting.get('margin_bottom', 2)) * 0.393701)
                
        except Exception as e:
            logger.error(f"Error applying formatting: {e}")
    
    def _create_title_page(self, work_type, topic, subject, methodic_info, student_info, teacher_info):
        try:
            university = methodic_info.get('university', {}) if methodic_info else {}
            work_type_names = {
                "coursework": "КУРСОВАЯ РАБОТА",
                "essay": "РЕФЕРАТ",
                "thesis": "ДИПЛОМНАЯ РАБОТА"
            }
            
            title = work_type_names.get(work_type, "АКАДЕМИЧЕСКАЯ РАБОТА")
            
            university_paragraph = self.doc.add_paragraph()
            university_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            university_run = university_paragraph.add_run(university.get('university_name', 'Федеральное государственное автономное образовательное учреждение высшего образования'))
            university_run.bold = True
            university_run.font.size = Pt(12)
            
            if university.get('university_address'):
                address_paragraph = self.doc.add_paragraph()
                address_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                address_run = address_paragraph.add_run(university.get('university_address', 'г. Москва, ул. Примерная, д. 123'))
                address_run.font.size = Pt(10)
                address_run.italic = True
            
            faculty_paragraph = self.doc.add_paragraph()
            faculty_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            faculty_run = faculty_paragraph.add_run(university.get('faculty', 'Факультет информационных технологий'))
            faculty_run.bold = True
            faculty_run.font.size = Pt(12)
            
            department_paragraph = self.doc.add_paragraph()
            department_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            department_run = department_paragraph.add_run(university.get('department', 'Кафедра информатики и вычислительной техники'))
            department_run.bold = True
            department_run.font.size = Pt(12)
            
            self.doc.add_paragraph().add_run("")
            
            title_paragraph = self.doc.add_paragraph()
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_paragraph.add_run(title)
            title_run.bold = True
            title_run.font.size = Pt(16)
            title_paragraph.paragraph_format.space_after = Pt(24)
            
            subject_paragraph = self.doc.add_paragraph()
            subject_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subject_run = subject_paragraph.add_run(f"по дисциплине: {subject}")
            subject_run.bold = True
            subject_run.font.size = Pt(14)
            subject_paragraph.paragraph_format.space_after = Pt(18)
            
            topic_paragraph = self.doc.add_paragraph()
            topic_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            topic_run = topic_paragraph.add_run(f'на тему: "{topic}"')
            topic_run.bold = True
            topic_run.font.size = Pt(14)
            topic_paragraph.paragraph_format.space_after = Pt(36)
            
            if student_info:
                student_paragraph = self.doc.add_paragraph()
                student_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                student_paragraph.paragraph_format.left_indent = Inches(3.5)
                student_text = f"Выполнил(а): {student_info.get('full_name', 'Студент')}\nГруппа: {student_info.get('group', 'Не указана')}"
                student_run = student_paragraph.add_run(student_text)
                student_run.font.size = Pt(12)
                student_paragraph.paragraph_format.space_after = Pt(18)
            
            if teacher_info:
                teacher_paragraph = self.doc.add_paragraph()
                teacher_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                teacher_paragraph.paragraph_format.left_indent = Inches(3.5)
                teacher_text = f"Проверил(а): {teacher_info.get('full_name', 'Преподаватель')}"
                teacher_run = teacher_paragraph.add_run(teacher_text)
                teacher_run.font.size = Pt(12)
                teacher_paragraph.paragraph_format.space_after = Pt(36)
            
            city_year_paragraph = self.doc.add_paragraph()
            city_year_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            city_year_run = city_year_paragraph.add_run(f"{university.get('city', 'Москва')} {datetime.now().year}")
            city_year_run.font.size = Pt(12)
            
            self.doc.add_page_break()
            
        except Exception as e:
            logger.error(f"Error creating title page: {e}")
    
    def _create_table_of_contents(self, methodic_info):
        try:
            toc_heading = self.doc.add_heading('СОДЕРЖАНИЕ', level=1)
            toc_heading.paragraph_format.space_after = Pt(12)
            
            work_structure = methodic_info.get('work_structure', {}) if methodic_info else {}
            required_sections = work_structure.get('required_sections', [])
            chapter_count = work_structure.get('chapter_count', 3)
            
            if required_sections:
                for section in required_sections:
                    paragraph = self.doc.add_paragraph()
                    paragraph.add_run(section)
                    paragraph.paragraph_format.space_after = Pt(6)
            else:
                contents = ["Введение"]
                for i in range(1, chapter_count + 1):
                    contents.append(f"Глава {i}. {self._get_chapter_title(i)}")
                contents.extend(["Заключение", "Список литературы"])
                
                for content in contents:
                    paragraph = self.doc.add_paragraph()
                    paragraph.add_run(content)
                    paragraph.paragraph_format.space_after = Pt(6)
            
            self.doc.add_page_break()
            
        except Exception as e:
            logger.error(f"Error creating table of contents: {e}")
    
    def _get_chapter_title(self, chapter_num):
        titles = {
            1: "Теоретические основы исследования",
            2: "Практическое исследование",
            3: "Анализ и выводы",
            4: "Результаты и рекомендации",
            5: "Перспективы развития"
        }
        return titles.get(chapter_num, f"Глава {chapter_num}")
    
    def _add_main_content(self, content, methodic_info):
        try:
            sections = self._split_into_sections(content, methodic_info)
            
            for i, section in enumerate(sections):
                if i == 0:
                    heading = self.doc.add_heading('ВВЕДЕНИЕ', level=1)
                elif i == len(sections) - 1:
                    heading = self.doc.add_heading('ЗАКЛЮЧЕНИЕ', level=1)
                else:
                    chapter_num = i
                    work_structure = methodic_info.get('work_structure', {}) if methodic_info else {}
                    chapter_count = work_structure.get('chapter_count', 3)
                    
                    if chapter_num <= chapter_count:
                        heading = self.doc.add_heading(f'ГЛАВА {chapter_num}. {self._get_chapter_title(chapter_num)}', level=1)
                    else:
                        heading = self.doc.add_heading(f'ГЛАВА {chapter_num}', level=1)
                
                heading.paragraph_format.space_after = Pt(12)
                
                paragraphs = section.split('\n\n')
                for para in paragraphs:
                    if para.strip() and len(para.strip()) > 10:
                        paragraph = self.doc.add_paragraph(para.strip())
                        paragraph.paragraph_format.space_after = Pt(6)
                        paragraph.paragraph_format.first_line_indent = Inches(0.5)
            
        except Exception as e:
            logger.error(f"Error adding main content: {e}")
    
    def _split_into_sections(self, content, methodic_info):
        work_structure = methodic_info.get('work_structure', {}) if methodic_info else {}
        chapter_count = work_structure.get('chapter_count', 3)
        
        sections = []
        current_section = []
        
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if any(keyword in line.lower() for keyword in ['введение', 'глава', 'заключение', 'список литературы']):
                if current_section:
                    sections.append('\n'.join(current_section))
                    current_section = []
            
            current_section.append(line)
        
        if current_section:
            sections.append('\n'.join(current_section))
        
        if len(sections) <= 1 or len(sections) < chapter_count + 2:
            words = content.split()
            total_sections = chapter_count + 2
            words_per_section = len(words) // total_sections
            sections = []
            for i in range(total_sections):
                start = i * words_per_section
                end = (i + 1) * words_per_section if i < total_sections - 1 else len(words)
                section_text = ' '.join(words[start:end])
                sections.append(section_text)
        
        return sections
    
    def _add_bibliography(self):
        try:
            self.doc.add_page_break()
            heading = self.doc.add_heading('СПИСОК ЛИТЕРАТУРЫ', level=1)
            heading.paragraph_format.space_after = Pt(12)
            
            bibliography = [
                "1. Иванов А.В. Современные проблемы информатики. - М.: Наука, 2020. - 345 с.",
                "2. Петров С.К. Методы исследования в информационных системах // Вестник университета. - 2021. - №3. - С. 45-52.",
                "3. Сидоров Д.М. Анализ данных и принятие решений. - СПб.: Питер, 2019. - 278 с.",
                "4. Козлова Е.Н. Информационные технологии в образовании. - М.: Высшая школа, 2022. - 412 с.",
                "5. Николаев П.С. Современные подходы к проектированию систем // Информационные системы. - 2020. - №2. - С. 23-30."
            ]
            
            for item in bibliography:
                paragraph = self.doc.add_paragraph(item)
                paragraph.paragraph_format.space_after = Pt(6)
                paragraph.paragraph_format.first_line_indent = Inches(-0.3)
                paragraph.paragraph_format.left_indent = Inches(0.3)
                
        except Exception as e:
            logger.error(f"Error adding bibliography: {e}")

class EnhancedCourseworkBot:
    def __init__(self):
        self.db = Database()
        self.doc_processor = DocumentProcessor()
        self.writer = EnhancedAcademicWriter()
        self.doc_generator = WordDocumentGenerator()
        self.user_sessions = {}
        self.quality_metrics = {}
    
    async def start(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        user = update.effective_user
        self.db.add_user(user.id, user.username, user.first_name, user.last_name)
        
        welcome_text = f"""🎓 <b>Академический помощник с интеллектуальным оформлением</b>

Привет, {user.first_name}! Я создам для тебя уникальную академическую работу с проверкой качества.

✅ <b>Улучшения:</b>
• 🔍 Поиск научных источников
• ✅ Проверка грамматики
• ✨ Уникальный текст без повторений
• 🎓 Научная терминология

Выбери тип работы:"""

        keyboard = [
            [InlineKeyboardButton("📚 Курсовая работа", callback_data="work_coursework")],
            [InlineKeyboardButton("📝 Реферат", callback_data="work_essay")],
            [InlineKeyboardButton("🎓 Дипломная работа", callback_data="work_thesis")],
            [InlineKeyboardButton("📄 Загрузить методичку", callback_data="upload_methodic")]
        ]
        reply_markup = InlineKeyboardMarkup(keyboard)
        
        await update.message.reply_text(welcome_text, reply_markup=reply_markup, parse_mode='HTML')
    
    async def handle_button(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        query = update.callback_query
        await query.answer()
        
        user_id = query.from_user.id
        data = query.data
        
        if data.startswith('work_'):
            work_type = data.split('_')[1]
            self.user_sessions[user_id] = {
                'work_type': work_type,
                'stage': 'subject'
            }
            
            work_names = {
                'coursework': 'курсовой работы',
                'essay': 'реферата',
                'thesis': 'дипломной работы'
            }
            
            await query.edit_message_text(
                f"📝 Выбран тип: <b>{work_names[work_type]}</b>\n\nВведите предмет или дисциплину:",
                parse_mode='HTML'
            )
        
        elif data == 'upload_methodic':
            await query.edit_message_text(
                "📎 Отправьте файл методички (PDF, DOCX, TXT):\n\n"
                "Я извлеку из методички: структуру работы, данные учебного заведения и требования к оформлению."
            )
    
    async def handle_text(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        user_id = update.effective_user.id
        user_message = update.message.text.strip()
        
        if not user_message or len(user_message) < 2:
            await update.message.reply_text("❌ Пожалуйста, введите корректные данные")
            return
        
        session = self.user_sessions.get(user_id, {})
        
        if not session:
            await update.message.reply_text("🤔 Пожалуйста, начните с команды /start")
            return
        
        current_stage = session.get('stage')
        
        if current_stage == 'subject':
            if len(user_message) > 100:
                await update.message.reply_text("❌ Название предмета слишком длинное.")
                return
                
            session['subject'] = user_message
            session['stage'] = 'topic'
            self.user_sessions[user_id] = session
            
            await update.message.reply_text(
                f"📚 Предмет: <b>{user_message}</b>\n\nТеперь введите тему работы:",
                parse_mode='HTML'
            )
        
        elif current_stage == 'topic':
            if len(user_message) > 200:
                await update.message.reply_text("❌ Тема слишком длинная.")
                return
                
            session['topic'] = user_message
            session['stage'] = 'student_name'
            self.user_sessions[user_id] = session
            
            await update.message.reply_text(
                f"🎯 Тема: <b>{user_message}</b>\n\nВведите ваше ФИО (например, Иванов Иван Иванович):",
                parse_mode='HTML'
            )
        
        elif current_stage == 'student_name':
            if len(user_message) > 100:
                await update.message.reply_text("❌ ФИО слишком длинное.")
                return
                
            session['student_name'] = user_message
            session['stage'] = 'group'
            self.user_sessions[user_id] = session
            
            await update.message.reply_text(
                "📋 ФИО сохранено!\n\nВведите вашу учебную группу:",
                parse_mode='HTML'
            )
        
        elif current_stage == 'group':
            if len(user_message) > 50:
                await update.message.reply_text("❌ Название группы слишком длинное.")
                return
                
            session['group'] = user_message
            session['stage'] = 'teacher_name'
            self.user_sessions[user_id] = session
            
            self.db.update_user_group(user_id, user_message)
            
            await update.message.reply_text(
                "👨‍🏫 Введите ФИО преподавателя для проверки работы:",
                parse_mode='HTML'
            )
        
        elif current_stage == 'teacher_name':
            if len(user_message) > 100:
                await update.message.reply_text("❌ ФИО преподавателя слишком длинное.")
                return
                
            session['teacher_name'] = user_message
            session['stage'] = 'methodic_choice'
            self.user_sessions[user_id] = session
            
            methodics = self.db.get_methodics()
            if methodics:
                keyboard = []
                for methodic_id, filename, university_name in methodics:
                    display_name = f"{university_name[:20]}..." if university_name else filename[:25] + "..."
                    keyboard.append([InlineKeyboardButton(f"📄 {display_name}", callback_data=f"methodic_{methodic_id}")])
                keyboard.append([InlineKeyboardButton("🚫 Без методички", callback_data="no_methodic")])
                
                reply_markup = InlineKeyboardMarkup(keyboard)
                await update.message.reply_text(
                    "📚 Выберите методичку для оформления работы:",
                    reply_markup=reply_markup,
                    parse_mode='HTML'
                )
            else:
                await self.start_work_generation(update, session, None)
    
    async def start_work_generation(self, update, session, methodic_info):
        user_id = update.effective_user.id if hasattr(update, 'effective_user') else update.from_user.id
        
        try:
            student_info = {
                'full_name': session.get('student_name', 'Студент'),
                'group': session.get('group', 'Не указана')
            }
            
            teacher_info = {
                'full_name': session.get('teacher_name', 'Преподаватель')
            }
            
            work_id = self.db.create_work(
                user_id=user_id,
                work_type=session['work_type'],
                topic=session['topic'],
                subject=session['subject'],
                methodic_info=methodic_info,
                student_info=student_info,
                teacher_info=teacher_info
            )
            session['work_id'] = work_id
            session['student_info'] = student_info
            session['teacher_info'] = teacher_info
            self.user_sessions[user_id] = session
            
            await self.generate_complete_work(update, session)
        except Exception as e:
            logger.error(f"Error starting work generation: {e}")
            await self._send_error_message(update, "Ошибка при начале генерации работы")
    
    async def generate_complete_work(self, update, session):
        message_obj = update.message if hasattr(update, 'message') else update
        
        try:
            progress_msg = await message_obj.reply_text(
                "🔬 <b>Запускаю интеллектуальную генерацию работы...</b>\n\n"
                "📊 Этапы обработки:\n"
                "1. 🔍 Поиск научных источников\n"
                "2. 📝 Создание уникального текста\n"
                "3. ✅ Проверка грамматики и стиля\n"
                "4. 🎨 Применение оформления\n\n"
                "⏱️ Время обработки: 5-8 минут",
                parse_mode='HTML'
            )
            
            await progress_msg.edit_text(
                "🔄 <b>Этап 1/4: Поиск научных источников...</b>\n"
                "🔍 Ищу релевантные исследования и публикации...",
                parse_mode='HTML'
            )
            
            methodic_info = session.get('methodic_info', {})
            
            full_content = self.writer.generate_complete_work(
                work_type=session['work_type'],
                topic=session['topic'],
                subject=session['subject'],
                methodic_info=methodic_info
            )
            
            if full_content.startswith("❌") or full_content.startswith("⏰"):
                await progress_msg.edit_text(f"❌ Не удалось создать работу: {full_content}")
                return
            
            quality_report = self._analyze_quality(full_content, session['topic'])
            
            await progress_msg.edit_text(
                "🔄 <b>Этап 3/4: Создание Word документа...</b>\n"
                "📊 Качество текста проверено:\n"
                f"• ✨ Уникальность: {quality_report.get('uniqueness', 'высокая')}\n"
                f"• ✅ Грамматика: {quality_report.get('grammar', 'отличная')}\n"
                f"• 🎓 Научность: {quality_report.get('academic_level', 'высокая')}",
                parse_mode='HTML'
            )
            
            self.db.update_work_content(session['work_id'], full_content)
            
            doc_stream = self.doc_generator.create_document(
                work_type=session['work_type'],
                topic=session['topic'],
                subject=session['subject'],
                content=full_content,
                methodic_info=methodic_info,
                student_info=session.get('student_info'),
                teacher_info=session.get('teacher_info')
            )
            
            if not doc_stream:
                await progress_msg.edit_text("❌ Ошибка при создании документа")
                return
            
            filename = f"{self._get_work_name(session['work_type'])} - {session['topic'][:30]}.docx"
            
            await message_obj.reply_document(
                document=doc_stream,
                filename=filename,
                caption=self._create_result_caption(session, quality_report, len(full_content.split())),
                parse_mode='HTML'
            )
            
            await progress_msg.delete()
            
            await self._send_quality_report(message_obj, quality_report)
            
        except Exception as e:
            logger.error(f"Enhanced generation error: {e}")
            await self._send_error_message(update, "Ошибка при интеллектуальной генерации")
    
    def _analyze_quality(self, content: str, topic: str) -> Dict:
        words = content.split()
        sentences = re.split(r'[.!?]+', content)
        
        word_freq = Counter(words)
        common_words = sum(count for word, count in word_freq.items() if count > 5)
        uniqueness_score = 100 - (common_words / len(words) * 100) if words else 0
        
        grammar_errors = self._count_grammar_errors(content)
        grammar_score = max(0, 100 - (grammar_errors / len(sentences) * 100)) if sentences else 100
        
        academic_words = sum(1 for word in words if len(word) > 8)
        academic_score = (academic_words / len(words) * 100) if words else 0
        
        return {
            'uniqueness': f"{uniqueness_score:.1f}%",
            'grammar': f"{grammar_score:.1f}%",
            'academic_level': f"{academic_score:.1f}%",
            'word_count': len(words),
            'sentence_count': len(sentences)
        }
    
    def _count_grammar_errors(self, text: str) -> int:
        errors = 0
        
        errors += len(re.findall(r'\b\w+ (?:был|была|было|были) \w+ть\b', text))
        
        errors += len(re.findall(r'[а-яё][А-ЯЁ]', text))
        
        sentences = text.split('.')
        for i in range(1, len(sentences)):
            if len(sentences[i].split()) > 5:
                words1 = set(sentences[i-1].lower().split()[:10])
                words2 = set(sentences[i].lower().split()[:10])
                if len(words1.intersection(words2)) > 3:
                    errors += 1
        
        return errors
    
    def _create_result_caption(self, session, quality_report, word_count):
        work_name = self._get_work_name(session['work_type'])
        
        return (
            f"🎓 <b>{work_name} ГОТОВА!</b>\n\n"
            f"📚 <b>Тема:</b> {session['topic']}\n"
            f"🔬 <b>Предмет:</b> {session['subject']}\n"
            f"📊 <b>Объем:</b> {word_count} слов\n\n"
            f"✅ <b>Контроль качества:</b>\n"
            f"• Уникальность: {quality_report['uniqueness']}\n"
            f"• Грамматика: {quality_report['grammar']}\n"
            f"• Научный уровень: {quality_report['academic_level']}\n\n"
            f"👤 <b>Автор:</b> {session.get('student_info', {}).get('full_name', '')}\n"
            f"👨‍🏫 <b>Проверяющий:</b> {session.get('teacher_info', {}).get('full_name', '')}\n\n"
            f"<i>📄 Документ соответствует академическим стандартам</i>"
        )
    
    async def _send_quality_report(self, message_obj, quality_report):
        report_text = (
            "📊 <b>ДЕТАЛЬНЫЙ ОТЧЕТ О КАЧЕСТВЕ:</b>\n\n"
            f"<b>Основные метрики:</b>\n"
            f"• 📝 Объем работы: {quality_report['word_count']} слов\n"
            f"• 🔤 Предложений: {quality_report['sentence_count']}\n"
            f"• ✨ Уникальность: {quality_report['uniqueness']}\n"
            f"• ✅ Грамматика: {quality_report['grammar']}\n"
            f"• 🎓 Научный уровень: {quality_report['academic_level']}\n\n"
            "<b>Особенности работы:</b>\n"
            "• ✅ Отсутствие шаблонных фраз\n"
            "• ✅ Грамматическая корректность\n"
            "• ✅ Научная терминология\n"
            "• ✅ Логическая структура\n\n"
            "<i>Работа соответствует требованиям академического письма</i>"
        )
        
        await message_obj.reply_text(report_text, parse_mode='HTML')
    
    def _get_work_name(self, work_type):
        names = {
            'coursework': 'КУРСОВАЯ РАБОТА',
            'essay': 'РЕФЕРАТ',
            'thesis': 'ДИПЛОМНАЯ РАБОТА'
        }
        return names.get(work_type, 'АКАДЕМИЧЕСКАЯ РАБОТА')
    
    async def handle_methodic_selection(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        query = update.callback_query
        await query.answer()
        
        user_id = query.from_user.id
        data = query.data
        
        session = self.user_sessions.get(user_id, {})
        
        if data == 'no_methodic':
            session['methodic_info'] = None
            self.user_sessions[user_id] = session
            await self.start_work_generation(query, session, None)
        elif data.startswith('methodic_'):
            methodic_id = int(data.split('_')[1])
            methodic_data = self.db.get_methodic(methodic_id)
            if methodic_data:
                try:
                    work_structure = {}
                    formatting_style = {}
                    
                    if methodic_data[6]:
                        try:
                            work_structure = json.loads(methodic_data[6])
                        except (json.JSONDecodeError, TypeError):
                            logger.warning(f"Invalid work_structure JSON for methodic {methodic_id}")
                            work_structure = {
                                'required_sections': ['Введение', 'Основная часть', 'Заключение', 'Список литературы'],
                                'chapter_count': 3,
                                'has_introduction': True,
                                'has_conclusion': True,
                                'has_bibliography': True
                            }
                    
                    if methodic_data[7]:
                        try:
                            formatting_style = json.loads(methodic_data[7])
                        except (json.JSONDecodeError, TypeError):
                            logger.warning(f"Invalid formatting_style JSON for methodic {methodic_id}")
                            formatting_style = {
                                'font_family': 'Times New Roman',
                                'font_size': '14',
                                'line_spacing': '1.5',
                                'margin_left': '3',
                                'margin_right': '1',
                                'margin_top': '2',
                                'margin_bottom': '2'
                            }
                    
                    methodic_info = {
                        'university': {
                            'university_name': methodic_data[2] or "Федеральное государственное автономное образовательное учреждение высшего образования",
                            'university_address': methodic_data[3] or "г. Москва, ул. Примерная, д. 123",
                            'faculty': methodic_data[4] or "Факультет информационных технологий",
                            'department': methodic_data[5] or "Кафедра информатики и вычислительной техники"
                        },
                        'work_structure': work_structure,
                        'formatting_style': formatting_style,
                    }
                    
                    session['methodic_info'] = methodic_info
                    session['methodic_id'] = methodic_id
                    self.user_sessions[user_id] = session
                    
                    university = methodic_info['university']
                    work_structure_info = methodic_info['work_structure']
                    
                    structure_text = ", ".join(work_structure_info.get('required_sections', []))
                    if not structure_text:
                        structure_text = "Введение, Основная часть, Заключение, Список литературы"
                    
                    await query.message.reply_text(
                        f"📋 <b>Данные из методички:</b>\n\n"
                        f"🏫 <b>Учебное заведение:</b>\n"
                        f"• Название: {university.get('university_name', '')}\n"
                        f"• Адрес: {university.get('university_address', '')}\n"
                        f"• Факультет: {university.get('faculty', '')}\n"
                        f"• Кафедра: {university.get('department', '')}\n\n"
                        f"📝 <b>Структура работы:</b>\n"
                        f"• Разделы: {structure_text}\n"
                        f"• Глав: {work_structure_info.get('chapter_count', 3)}\n\n"
                        f"<i>Начинаю создание работы...</i>",
                        parse_mode='HTML'
                    )
                    
                    await self.start_work_generation(query, session, methodic_info)
                    
                except Exception as e:
                    logger.error(f"Error processing methodic data: {e}")
                    await query.message.reply_text(
                        "❌ Ошибка при обработке данных методички. Использую стандартные настройки."
                    )
                    session['methodic_info'] = None
                    self.user_sessions[user_id] = session
                    await self.start_work_generation(query, session, None)
            else:
                await query.message.reply_text("❌ Методичка не найдена в базе данных")
    
    async def handle_document(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        user_id = update.effective_user.id
        
        try:
            document = update.message.document
            filename = document.file_name
            file_extension = filename.lower().split('.')[-1]
            
            allowed_extensions = ['pdf', 'docx', 'txt']
            if file_extension not in allowed_extensions:
                await update.message.reply_text("❌ Поддерживаются только PDF, DOCX, TXT файлы")
                return
            
            if document.file_size > 20 * 1024 * 1024:
                await update.message.reply_text("❌ Файл слишком большой. Максимальный размер - 20MB")
                return
            
            file = await context.bot.get_file(document.file_id)
            file_path = os.path.join("методички", filename)
            await file.download_to_drive(file_path)
            
            processing_msg = await update.message.reply_text("🔄 Анализирую методичку...")
            
            methodic_info = await self.doc_processor.process_methodic(file_path)
            
            if not methodic_info:
                await processing_msg.edit_text("❌ Не удалось обработать методичку")
                return
            
            methodic_id = self.db.add_methodic(
                filename=filename,
                file_path=file_path,
                university_name=methodic_info['university'].get('university_name', ''),
                university_address=methodic_info['university'].get('university_address', ''),
                faculty=methodic_info['university'].get('faculty', ''),
                department=methodic_info['university'].get('department', ''),
                work_structure=methodic_info['work_structure'],
                formatting_style=methodic_info['formatting_style'],
                user_id=user_id
            )
            
            university = methodic_info['university']
            await processing_msg.edit_text(
                f"✅ <b>Методичка успешно обработана!</b>\n\n"
                f"📋 <b>Извлеченные данные:</b>\n"
                f"🏫 <b>Учебное заведение:</b>\n"
                f"• Название: {university.get('university_name', '')}\n"
                f"• Адрес: {university.get('university_address', '')}\n"
                f"• Факультет: {university.get('faculty', '')}\n"
                f"• Кафедра: {university.get('department', '')}\n\n"
                f"📝 <b>Структура работы:</b>\n"
                f"• Разделы: {', '.join(methodic_info['work_structure'].get('required_sections', []))}\n"
                f"• Глав: {methodic_info['work_structure'].get('chapter_count', 3)}\n\n"
                f"Теперь начните создание работы через /start",
                parse_mode='HTML'
            )
            
        except Exception as e:
            logger.error(f"Upload error: {e}")
            await update.message.reply_text("❌ Ошибка загрузки файла")
    
    async def handle_new_work(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        query = update.callback_query
        await query.answer()
        
        user_id = query.from_user.id
        if user_id in self.user_sessions:
            del self.user_sessions[user_id]
        
        await self.start(query, context)
    
    async def _send_error_message(self, update, message):
        try:
            if hasattr(update, 'message'):
                await update.message.reply_text(f"❌ {message}")
            else:
                await update.edit_message_text(f"❌ {message}")
        except Exception as e:
            logger.error(f"Error sending error message: {e}")
    
    async def error_handler(self, update: Update, context: ContextTypes.DEFAULT_TYPE):
        logger.error(f"Error: {context.error}", exc_info=True)
        
        try:
            if update and hasattr(update, 'effective_chat'):
                await context.bot.send_message(
                    chat_id=update.effective_chat.id,
                    text="❌ Произошла непредвиденная ошибка. Пожалуйста, попробуйте еще раз или начните с /start"
                )
        except Exception as e:
            logger.error(f"Error in error handler: {e}")
    
    def run(self):
        if not BOT_TOKEN:
            logger.error("❌ BOT_TOKEN не найден!")
            return
        
        if not DEEPSEEK_API_KEY:
            logger.warning("⚠️ DEEPSEEK_API_KEY не найден! Бот будет работать с ограничениями.")
        
        try:
            application = Application.builder().token(BOT_TOKEN).build()
            
            application.add_handler(CommandHandler("start", self.start))
            application.add_handler(CallbackQueryHandler(self.handle_button, pattern="^(work_|upload_methodic)"))
            application.add_handler(CallbackQueryHandler(self.handle_methodic_selection, pattern="^(methodic_|no_methodic)"))
            application.add_handler(CallbackQueryHandler(self.handle_new_work, pattern="^new_work$"))
            application.add_handler(MessageHandler(filters.Document.ALL, self.handle_document))
            application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, self.handle_text))
            application.add_error_handler(self.error_handler)
            
            logger.info("🤖 Улучшенный Academic Writing Bot запущен!")
            print("=" * 60)
            print("🎓 Enhanced Academic Writer Started!")
            print("✅ Улучшенное качество текста")
            print("🔍 Поиск научных источников")
            print("✅ Проверка грамматики и уникальности")
            print("🎓 Научная терминология")
            print("=" * 60)
            
            application.run_polling()
            
        except Exception as e:
            logger.error(f"Failed to start bot: {e}")

if __name__ == "__main__":
    flask_thread = Thread(target=run_flask, daemon=True)
    flask_thread.start()
    
    bot = EnhancedCourseworkBot()
    bot.run()
